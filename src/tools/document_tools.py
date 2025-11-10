"""
Document Processing Tools
=========================

Tools for extracting and analyzing text from documents (PDFs, TXT, DOCX).

Use cases:
- Manual situation reports (SITREPs)
- Intelligence documents
- Mission orders
- After-action reports
- Any text-based tactical information

Capabilities:
- PDF text extraction
- Plain text file reading
- DOCX document parsing
- Text preprocessing and cleaning
- Structured output for entity extraction
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime, timezone

from langsmith import traceable


# ==================== PDF PROCESSING ====================

def _extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """
    Extract text from PDF file using PyPDF2
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Extracted text or None if error
    """
    try:
        import PyPDF2
        
        text_content = []
        
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():  # Only add non-empty pages
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
        
        return "\n\n".join(text_content)
        
    except ImportError:
        raise ImportError(
            "PyPDF2 not installed. Install with: pip install PyPDF2"
        )
    except Exception as e:
        print(f"❌ PDF extraction error: {e}")
        return None


# ==================== TEXT FILE PROCESSING ====================

def _read_text_file(text_path: str, encoding: str = 'utf-8') -> Optional[str]:
    """
    Read plain text file with fallback encodings
    
    Args:
        text_path: Path to text file
        encoding: Primary encoding to try
        
    Returns:
        File contents or None if error
    """
    # Try multiple encodings in order of preference
    encodings = [encoding, 'utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for enc in encodings:
        try:
            with open(text_path, 'r', encoding=enc) as file:
                return file.read()
        except (UnicodeDecodeError, LookupError):
            continue
        except Exception as e:
            print(f"❌ Error reading file with {enc}: {e}")
            return None
    
    print(f"❌ Failed to read {text_path} with any encoding")
    return None


# ==================== DOCX PROCESSING ====================

def _extract_text_from_docx(docx_path: str) -> Optional[str]:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text or None if error
    """
    try:
        import docx
        
        doc = docx.Document(docx_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables (if any)
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        # Combine all content
        full_text = "\n\n".join(paragraphs)
        
        if tables_text:
            full_text += "\n\n--- TABLES ---\n" + "\n".join(tables_text)
        
        return full_text
        
    except ImportError:
        raise ImportError(
            "python-docx not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        print(f"❌ DOCX extraction error: {e}")
        return None


# ==================== VALIDATION ====================

def validate_document_file(document_path: str) -> tuple[bool, Optional[str]]:
    """
    Validate document file exists and is supported
    
    Args:
        document_path: Path to check
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    # Check existence
    if not os.path.exists(document_path):
        return False, f"Document file not found: {document_path}"
    
    # Check format
    supported_extensions = {'.txt', '.pdf', '.doc', '.docx'}
    extension = Path(document_path).suffix.lower()
    
    if extension not in supported_extensions:
        return False, f"Unsupported document format: {extension}. Supported: {supported_extensions}"
    
    # Check file size (reasonable limit for text documents)
    file_size_mb = os.path.getsize(document_path) / (1024 * 1024)
    if file_size_mb > 50:  # 50MB is generous for text
        return False, f"Document too large: {file_size_mb:.1f}MB (max 50MB)"
    
    return True, None


# ==================== TEXT PREPROCESSING ====================

def clean_extracted_text(text: str, max_lines: Optional[int] = None) -> str:
    """
    Clean and preprocess extracted text
    
    Args:
        text: Raw extracted text
        max_lines: Maximum lines to keep (None = keep all)
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Strip whitespace
        line = line.strip()
        
        # Skip empty lines (but preserve some structure)
        if not line:
            # Only add empty line if previous wasn't empty
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        
        # Remove excessive spaces within line
        line = ' '.join(line.split())
        
        cleaned_lines.append(line)
    
    # Limit lines if requested
    if max_lines and len(cleaned_lines) > max_lines:
        cleaned_lines = cleaned_lines[:max_lines]
        cleaned_lines.append(f"\n... [Truncated - total {len(lines)} lines] ...")
    
    return '\n'.join(cleaned_lines)


# ==================== CORE EXTRACTION FUNCTIONS ====================

@traceable(name="extract_text_from_document")
def extract_text_from_document(
    document_path: str,
    clean_text: bool = True,
    max_lines: Optional[int] = None
) -> Dict[str, Any]:
    """
    Extract text from document (auto-detects format)
    
    Args:
        document_path: Path to document file
        clean_text: Whether to clean/preprocess text
        max_lines: Maximum lines to extract (None = all)
        
    Returns:
        Dict with:
            - success: bool
            - text: str (extracted text)
            - format: str (pdf, txt, docx)
            - num_pages: Optional[int] (for PDFs)
            - num_lines: int
            - error: Optional[str]
    """
    try:
        # Validate file
        is_valid, error_msg = validate_document_file(document_path)
        if not is_valid:
            return {
                "success": False,
                "text": "",
                "format": None,
                "num_pages": None,
                "num_lines": 0,
                "error": error_msg
            }
        
        # Detect format
        extension = Path(document_path).suffix.lower()
        
        # Extract based on format
        print(f"📄 Extracting text from: {Path(document_path).name}")
        
        if extension == '.pdf':
            raw_text = _extract_text_from_pdf(document_path)
            doc_format = "pdf"
        elif extension == '.txt':
            raw_text = _read_text_file(document_path)
            doc_format = "txt"
        elif extension in ['.doc', '.docx']:
            if extension == '.doc':
                return {
                    "success": False,
                    "text": "",
                    "format": "doc",
                    "num_pages": None,
                    "num_lines": 0,
                    "error": "Legacy .doc format not supported. Please convert to .docx or .pdf"
                }
            raw_text = _extract_text_from_docx(document_path)
            doc_format = "docx"
        else:
            return {
                "success": False,
                "text": "",
                "format": None,
                "num_pages": None,
                "num_lines": 0,
                "error": f"Unsupported format: {extension}"
            }
        
        # Check extraction success
        if raw_text is None:
            return {
                "success": False,
                "text": "",
                "format": doc_format,
                "num_pages": None,
                "num_lines": 0,
                "error": f"Failed to extract text from {doc_format.upper()}"
            }
        
        # Clean text if requested
        if clean_text:
            final_text = clean_extracted_text(raw_text, max_lines)
        else:
            final_text = raw_text
        
        # Count lines
        num_lines = len(final_text.split('\n'))
        
        return {
            "success": True,
            "text": final_text,
            "format": doc_format,
            "num_pages": None,  # Could extract from PDF metadata if needed
            "num_lines": num_lines,
            "error": None
        }
        
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "format": None,
            "num_pages": None,
            "num_lines": 0,
            "error": f"Document extraction failed: {str(e)}"
        }


# ==================== HIGH-LEVEL INTERFACE ====================

@traceable(name="process_document")
def process_document(
    document_path: str,
    max_lines: Optional[int] = 1000
) -> str:
    """
    High-level document processing function
    
    This is the main function that nodes should call.
    Automatically extracts text from any supported document format.
    
    Args:
        document_path: Path to document file
        max_lines: Maximum lines to extract (prevents huge docs)
        
    Returns:
        Formatted string report suitable for LLM consumption
        
    Example return:
        ```
        DOCUMENT EXTRACTION REPORT
        ==========================
        File: sitrep_20251015.pdf
        Format: PDF
        Lines extracted: 145
        Status: SUCCESS
        
        CONTENT:
        --------
        [Extracted text here...]
        
        ==========================
        ```
    """
    try:
        # Extract text
        result = extract_text_from_document(
            document_path=document_path,
            clean_text=True,
            max_lines=max_lines
        )
        
        file_name = Path(document_path).name
        
        # Check for errors
        if not result["success"]:
            return f"""
DOCUMENT EXTRACTION REPORT
==========================
File: {file_name}
Status: FAILED

ERROR: {result['error']}
==========================
"""
        
        # Success
        doc_format = result["format"].upper()
        num_lines = result["num_lines"]
        content = result["text"]
        
        # Check if content is empty
        if not content.strip():
            return f"""
DOCUMENT EXTRACTION REPORT
==========================
File: {file_name}
Format: {doc_format}
Status: WARNING

ISSUE: Document appears empty or contains no extractable text
==========================
"""
        
        return f"""
DOCUMENT EXTRACTION REPORT
==========================
File: {file_name}
Format: {doc_format}
Lines extracted: {num_lines}
Status: SUCCESS

CONTENT:
--------
{content}

==========================
"""
        
    except Exception as e:
        return f"""
DOCUMENT EXTRACTION REPORT
==========================
File: {Path(document_path).name}
Status: CRITICAL ERROR

ERROR: {str(e)}
==========================
"""


# ==================== UTILITY FUNCTIONS ====================

def is_document_file(file_path: str) -> bool:
    """
    Check if file is a supported document format
    
    Args:
        file_path: Path to check
        
    Returns:
        True if supported document file
    """
    document_extensions = {'.txt', '.pdf', '.doc', '.docx'}
    return Path(file_path).suffix.lower() in document_extensions


def get_document_info(document_path: str) -> Optional[Dict[str, Any]]:
    """
    Get basic document metadata without extracting content
    
    Args:
        document_path: Path to document
        
    Returns:
        Dict with metadata or None if error
    """
    try:
        if not os.path.exists(document_path):
            return None
        
        file_path = Path(document_path)
        file_size_mb = os.path.getsize(document_path) / (1024 * 1024)
        
        info = {
            "file_name": file_path.name,
            "file_extension": file_path.suffix.lower(),
            "file_size_mb": round(file_size_mb, 2),
            "file_modified": datetime.fromtimestamp(
                os.path.getmtime(document_path)
            ).isoformat()
        }
        
        # For PDFs, get page count
        if file_path.suffix.lower() == '.pdf':
            try:
                import PyPDF2
                with open(document_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    info["num_pages"] = len(reader.pages)
            except Exception:
                info["num_pages"] = None
        
        return info
        
    except Exception:
        return None


def extract_metadata_from_pdf(pdf_path: str) -> Optional[Dict[str, Any]]:
    """
    Extract metadata from PDF (author, title, creation date, etc.)
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Dict with metadata or None if error
    """
    try:
        import PyPDF2
        
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            metadata = reader.metadata
            
            if not metadata:
                return None
            
            # Extract common fields
            extracted = {}
            
            if metadata.get('/Title'):
                extracted['title'] = metadata['/Title']
            if metadata.get('/Author'):
                extracted['author'] = metadata['/Author']
            if metadata.get('/Subject'):
                extracted['subject'] = metadata['/Subject']
            if metadata.get('/Creator'):
                extracted['creator'] = metadata['/Creator']
            if metadata.get('/Producer'):
                extracted['producer'] = metadata['/Producer']
            if metadata.get('/CreationDate'):
                extracted['creation_date'] = metadata['/CreationDate']
            if metadata.get('/ModDate'):
                extracted['modification_date'] = metadata['/ModDate']
            
            return extracted if extracted else None
            
    except Exception:
        return None


def search_text_in_document(
    document_path: str,
    search_terms: List[str],
    case_sensitive: bool = False
) -> Dict[str, Any]:
    """
    Search for specific terms in document
    
    Args:
        document_path: Path to document
        search_terms: List of terms to search for
        case_sensitive: Whether search is case-sensitive
        
    Returns:
        Dict with search results:
            - success: bool
            - matches: Dict[str, List[int]] (term -> line numbers)
            - total_matches: int
            - error: Optional[str]
    """
    try:
        # Extract text
        result = extract_text_from_document(document_path, clean_text=False)
        
        if not result["success"]:
            return {
                "success": False,
                "matches": {},
                "total_matches": 0,
                "error": result["error"]
            }
        
        text = result["text"]
        lines = text.split('\n')
        
        # Search for each term
        matches = {}
        total_matches = 0
        
        for term in search_terms:
            term_matches = []
            search_term = term if case_sensitive else term.lower()
            
            for line_num, line in enumerate(lines, 1):
                search_line = line if case_sensitive else line.lower()
                
                if search_term in search_line:
                    term_matches.append(line_num)
                    total_matches += 1
            
            if term_matches:
                matches[term] = term_matches
        
        return {
            "success": True,
            "matches": matches,
            "total_matches": total_matches,
            "error": None
        }
        
    except Exception as e:
        return {
            "success": False,
            "matches": {},
            "total_matches": 0,
            "error": f"Search failed: {str(e)}"
        }